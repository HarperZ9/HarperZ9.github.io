"""Build deterministic PDF/DOCX career artifacts from reviewed HTML."""
from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import datetime, timezone
import hashlib
from html import escape
from html.parser import HTMLParser
from importlib.metadata import version
import json
import os
from pathlib import Path
import platform
import tempfile
import zipfile
from xml.etree import ElementTree

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from pypdf import PdfReader


SOURCE_EPOCH = 1788109200  # 2026-08-30T17:00:00Z


@dataclass(frozen=True, slots=True)
class ArtifactSpec:
    source: str
    stem: str
    lane_id: str


SPECS = (
    ArtifactSpec(
        "resume-support-operations.html",
        "Zain-Dana-Harper-Resume-Support-Developer-Operations-QA",
        "support-operations-qa",
    ),
    ArtifactSpec(
        "resume-evaluation-tooling.html",
        "Zain-Dana-Harper-Resume-Evaluation-Tooling-Python-Developer-Tools",
        "evaluation-python-tools",
    ),
    ArtifactSpec(
        "resume-public-operations.html",
        "Zain-Dana-Harper-Resume-Public-Operations",
        "public-operations",
    ),
    ArtifactSpec(
        "resume-grounds.html",
        "Zain-Dana-Harper-Resume-Grounds",
        "grounds",
    ),
    ArtifactSpec("cv.html", "Zain-Dana-Harper-CV", "page:cv"),
)


@dataclass(frozen=True, slots=True)
class Block:
    kind: str
    text: str


class _ArticleParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._article_depth = 0
        self._kind: str | None = None
        self._parts: list[str] = []
        self.blocks: list[Block] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag == "article":
            self._article_depth += 1
        elif self._article_depth and tag in {"h1", "h2", "p", "li"}:
            self._kind = tag
            self._parts = []
        elif (
            self._article_depth
            and self._kind == "p"
            and tag == "span"
            and any(part.strip() for part in self._parts)
        ):
            self._parts.append(" | ")

    def handle_data(self, data: str) -> None:
        if self._article_depth and self._kind is not None:
            self._parts.append(data)

    def handle_endtag(self, tag: str) -> None:
        if self._article_depth and tag == self._kind:
            text = " ".join("".join(self._parts).split())
            if text:
                self.blocks.append(Block(tag, text))
            self._kind = None
            self._parts = []
        if tag == "article" and self._article_depth:
            self._article_depth -= 1


class _InvariantCanvas(canvas.Canvas):
    def __init__(self, *args, title: str, **kwargs) -> None:
        kwargs["invariant"] = 1
        super().__init__(*args, **kwargs)
        self.setAuthor("Zain Dana Harper")
        self.setCreator("harperz9 career pipeline")
        self.setTitle(title)
        self.setSubject("Public career document")


def _read_blocks(path: Path) -> tuple[Block, ...]:
    parser = _ArticleParser()
    parser.feed(path.read_text(encoding="utf-8"))
    if not parser.blocks or parser.blocks[0].kind != "h1":
        raise ValueError(f"no career article heading in {path.name}")
    return tuple(parser.blocks)


def _build_pdf(path: Path, blocks: tuple[Block, ...]) -> None:
    title = blocks[0].text
    styles = {
        "h1": ParagraphStyle(
            "CareerTitle", fontName="Helvetica-Bold", fontSize=13,
            leading=14, spaceAfter=4, alignment=TA_LEFT,
        ),
        "h2": ParagraphStyle(
            "CareerHeading", fontName="Helvetica-Bold", fontSize=9.6,
            leading=10.5, spaceBefore=2.4, spaceAfter=0.8,
        ),
        "p": ParagraphStyle(
            "CareerBody", fontName="Helvetica", fontSize=8.5,
            leading=9.7, spaceAfter=0.9,
        ),
        "li": ParagraphStyle(
            "CareerBullet", fontName="Helvetica", fontSize=8.4,
            leading=9.6, leftIndent=8, firstLineIndent=-5, spaceAfter=0.7,
        ),
    }
    story = []
    for block in blocks:
        payload = escape(block.text)
        if block.kind == "li":
            payload = "&#8226; " + payload
        story.append(Paragraph(payload, styles[block.kind]))
    story.append(Spacer(1, 0.01 * inch))
    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER,
        leftMargin=0.48 * inch, rightMargin=0.48 * inch,
        topMargin=0.4 * inch, bottomMargin=0.4 * inch,
        pageCompression=1,
    )

    def canvasmaker(*args, **kwargs):
        return _InvariantCanvas(*args, title=title, **kwargs)

    doc.build(story, canvasmaker=canvasmaker)


def _repack_docx(path: Path, fixed_datetime: datetime) -> None:
    with tempfile.TemporaryDirectory() as temp_dir:
        temporary = Path(temp_dir) / path.name
        os.replace(path, temporary)
        with zipfile.ZipFile(temporary, "r") as source, zipfile.ZipFile(
            path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9,
        ) as target:
            for name in sorted(source.namelist()):
                info = zipfile.ZipInfo(name, fixed_datetime.timetuple()[:6])
                info.compress_type = zipfile.ZIP_DEFLATED
                info.create_system = 0
                info.external_attr = 0o600 << 16
                target.writestr(info, source.read(name))


def _build_docx(
    path: Path,
    blocks: tuple[Block, ...],
    fixed_datetime: datetime,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(9)
    for block in blocks:
        if block.kind == "h1":
            paragraph = document.add_paragraph(block.text, style="Title")
        elif block.kind == "h2":
            paragraph = document.add_paragraph(block.text, style="Heading 2")
        elif block.kind == "li":
            paragraph = document.add_paragraph(block.text, style="List Bullet")
        else:
            paragraph = document.add_paragraph(block.text)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1.5)
        paragraph.paragraph_format.line_spacing = 1.0

    properties = document.core_properties
    properties.title = blocks[0].text
    properties.subject = "Public career document"
    properties.author = "Zain Dana Harper"
    properties.last_modified_by = "harperz9 career pipeline"
    properties.created = fixed_datetime
    properties.modified = fixed_datetime
    properties.revision = 1
    document.save(path)
    _repack_docx(path, fixed_datetime)


def _sha256(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def _docx_text(path: Path) -> str:
    with zipfile.ZipFile(path) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return "\n".join(node.text or "" for node in root.iter(namespace + "t"))


def _artifact_text(path: Path) -> tuple[str, int | None]:
    if path.suffix == ".pdf":
        reader = PdfReader(path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text, len(reader.pages)
    return _docx_text(path), None


def _normalized_extraction(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.splitlines()).strip() + "\n"


def _build_receipt(
    site_root: Path,
    output_root: Path,
    source_epoch: int,
) -> dict:
    rows = []
    sources = []
    for spec in SPECS:
        source_payload = (site_root / spec.source).read_bytes()
        sources.append({"path": spec.source, "sha256": _sha256(source_payload)})
        for suffix, mime_type in (
            ("pdf", "application/pdf"),
            ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ):
            path = output_root / f"{spec.stem}.{suffix}"
            payload = path.read_bytes()
            text, page_count = _artifact_text(path)
            extraction = _normalized_extraction(text)
            rows.append({
                "byte_length": len(payload),
                "extraction_sha256": _sha256(extraction.encode("utf-8")),
                "format": suffix,
                "lane_id": spec.lane_id,
                "mime_type": mime_type,
                "page_count": page_count,
                "path": f"career/{path.name}",
                "sha256": _sha256(payload),
                "source": spec.source,
            })
    build_inputs = []
    for relative in (
        "tools/build_career_artifacts.py",
        "requirements-career-docs.txt",
    ):
        payload = (site_root / relative).read_bytes()
        build_inputs.append({"path": relative, "sha256": _sha256(payload)})
    dependencies = {
        name: version(name)
        for name in (
            "charset-normalizer",
            "lxml",
            "pillow",
            "PyMuPDF",
            "python-docx",
            "pypdf",
            "reportlab",
            "typing_extensions",
        )
    }
    return {
        "artifacts": sorted(rows, key=lambda row: row["path"]),
        "build_inputs": build_inputs,
        "runtime": {
            "dependencies": dependencies,
            "python": platform.python_version(),
        },
        "schema": "harperz9-career-build/v1",
        "source_epoch": source_epoch,
        "sources": sorted(sources, key=lambda row: row["path"]),
    }


def _current_html_target(site_root: Path, relative: str) -> Path:
    path = Path(relative)
    if (
        not relative
        or path.is_absolute()
        or path.drive
        or relative != path.as_posix()
        or any(part in {".", ".."} for part in path.parts)
    ):
        raise ValueError(
            "release manifest current_html paths must be normalized relative paths"
        )
    target = (site_root / path).resolve()
    if not target.is_relative_to(site_root):
        raise ValueError(
            "release manifest current_html paths must stay within the site root"
        )
    return target


def _validate_release_manifest(path: Path, site_root: Path) -> dict:
    manifest = json.loads(path.read_text(encoding="utf-8"))
    generated_paths = {
        f"career/{spec.stem}.{suffix}"
        for spec in SPECS
        for suffix in ("pdf", "docx")
    }
    artifact_paths = {row["path"] for row in manifest["artifacts"]}
    missing_artifacts = sorted(generated_paths - artifact_paths)
    if missing_artifacts:
        raise ValueError(
            f"release manifest has no rows for: {', '.join(missing_artifacts)}"
        )
    source_names = {spec.source for spec in SPECS}
    html_paths = {row["path"] for row in manifest["current_html"]}
    missing_html = sorted(source_names - html_paths)
    if missing_html:
        raise ValueError(
            f"release manifest has no HTML rows for: {', '.join(missing_html)}"
        )
    html_targets = {
        relative: _current_html_target(site_root, relative)
        for relative in html_paths
    }
    missing_targets = sorted(
        relative for relative, target in html_targets.items() if not target.is_file()
    )
    if missing_targets:
        raise ValueError(
            "release manifest HTML targets do not exist: "
            + ", ".join(missing_targets)
        )
    for relative in sorted(html_paths):
        _read_blocks(html_targets[relative])
    return manifest


def _write_json_atomic(path: Path, value: dict) -> None:
    payload = json.dumps(value, indent=2, sort_keys=True) + "\n"
    temporary_name: str | None = None
    try:
        with tempfile.NamedTemporaryFile(
            "w",
            encoding="utf-8",
            newline="\n",
            dir=path.parent,
            prefix=f".{path.name}.",
            suffix=".tmp",
            delete=False,
        ) as temporary:
            temporary.write(payload)
            temporary_name = temporary.name
        os.replace(temporary_name, path)
    finally:
        if temporary_name is not None:
            temporary_path = Path(temporary_name)
            if temporary_path.exists():
                temporary_path.unlink()


def _update_release_manifest(
    path: Path,
    manifest: dict,
    receipt: dict,
    source_epoch: int,
    site_root: Path,
) -> None:
    generated = {
        row["path"]: {key: value for key, value in row.items() if key != "source"}
        for row in receipt["artifacts"]
    }

    manifest["artifacts"] = [
        generated.get(row["path"], row) for row in manifest["artifacts"]
    ]
    manifest["generated_at_epoch"] = source_epoch
    updated = []
    for row in manifest["current_html"]:
        source_path = _current_html_target(site_root, row["path"])
        payload = source_path.read_bytes()
        extraction = "\n".join(
            block.text for block in _read_blocks(source_path)
        ).strip() + "\n"
        current = dict(row)
        current["byte_length"] = len(payload)
        current["sha256"] = _sha256(payload)
        current["extraction_sha256"] = _sha256(extraction.encode("utf-8"))
        updated.append(current)
    manifest["current_html"] = updated
    _write_json_atomic(path, manifest)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--site-root", type=Path, required=True)
    parser.add_argument("--output-root", type=Path, required=True)
    parser.add_argument("--receipt", type=Path, required=True)
    parser.add_argument("--source-epoch", type=int, default=SOURCE_EPOCH)
    parser.add_argument("--release-manifest", type=Path)
    args = parser.parse_args()

    site_root = args.site_root.resolve()
    release_manifest = None
    if args.release_manifest is not None:
        release_manifest = _validate_release_manifest(
            args.release_manifest,
            site_root,
        )
    blocks_by_spec = {
        spec: _read_blocks(site_root / spec.source) for spec in SPECS
    }
    fixed_datetime = datetime.fromtimestamp(args.source_epoch, timezone.utc)
    args.output_root.mkdir(parents=True, exist_ok=True)
    for spec in SPECS:
        blocks = blocks_by_spec[spec]
        _build_pdf(args.output_root / f"{spec.stem}.pdf", blocks)
        _build_docx(
            args.output_root / f"{spec.stem}.docx",
            blocks,
            fixed_datetime,
        )
    receipt = _build_receipt(site_root, args.output_root, args.source_epoch)
    args.receipt.parent.mkdir(parents=True, exist_ok=True)
    _write_json_atomic(args.receipt, receipt)
    if args.release_manifest is not None and release_manifest is not None:
        _update_release_manifest(
            args.release_manifest,
            release_manifest,
            receipt,
            args.source_epoch,
            site_root,
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
